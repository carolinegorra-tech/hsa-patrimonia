"""
checklist_builder.py — Gera o documento "Lista Inicial de Documentos" em .docx.

Versão LEVE (sem template embutido). Gera o documento do zero usando
python-docx. Não tem o logo da HSA como imagem — o cabeçalho é texto
formatado "HUMBERTO SANCHES + ASSOCIADOS". Esta versão é compatível com
o Vercel serverless (arquivo pequeno, sem deps pesadas no import).

Args do build_checklist:
  - client_name: nome completo do cliente
  - file_status: dict {key: bool} com as 12 keys
  - output_path: onde salvar o .docx
"""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


# 12 itens na ordem que aparecem no template original
_ITEMS = [
    ("certidao",
     'Cópia da certidão de casamento de {NAME} ("{FIRST}") e esposa, bem como de eventual pacto antenupcial entre eles, se aplicável.'),
    ("familia",
     'Informações relevantes sobre a estrutura familiar de {FIRST} e esposa, incluindo nome e idade dos filhos, netos e respectivos estados civis.'),
    ("dirpf",
     'Declaração(ões) de Imposto de Renda sobre a Pessoa Física ("DIRPF") de {FIRST} e esposa, ano-calendário 2024 e/ou rascunho da DIRPF referente ao ano-calendário 2025.'),
    ("dcbe",
     'Declaração(ões) de Capitais Brasileiros no Exterior ("DCBE") mais recente entregue ao Banco Central por {FIRST} e esposa, conforme aplicável.'),
    ("alteracoes",
     'Lista de alterações patrimoniais relevantes ocorridas no decorrer dos anos de 2025 e 2026 e/ou com previsão de materialização no curto/médio prazo, incluindo ativos com expectativa de recebimento.'),
    ("societario_br",
     'Se aplicável, organograma e documentos societários (contratos sociais, estatutos, atas), balanço patrimonial mais recente e documentos que indiquem as regras de governança de sociedades investidas por {FIRST} e esposa no Brasil, e os respectivos percentuais de participação em cada veículo. Favor incluir informações sobre as atividades desempenhadas por cada sociedade (i.e., se operacional (indicar ramo de atividade), holding patrimonial, inativa etc.).'),
    ("societario_off",
     'Organograma e documentos societários dos veículos offshore detidos por {FIRST} e esposa, incluindo sociedades, fundos e trusts, tais como Memorandum and Articles of Association, Register of Members, Register of Directors e Trust Deed, conforme aplicável.'),
    ("lei_14754",
     'Informações quanto ao tratamento tributário aplicável aos eventuais veículos controlados no exterior por {FIRST} e esposa, especificamente para fins da Lei nº 14.754/23, se aplicável.'),
    ("imoveis",
     'Informações sobre a destinação de uso e fluxo de rendimentos dos imóveis detidos por {FIRST} e esposa, conforme aplicável.'),
    ("emprestimos",
     'Informações sobre eventuais empréstimos concedidos ou tomados por {FIRST} e esposa, em vigor, conforme aplicável.'),
    ("doacoes",
     'Informações sobre eventuais doações ou heranças recebidas ou realizadas por {FIRST} e esposa, incluindo informações sobre o recolhimento do ITCMD aplicável, conforme aplicável.'),
    ("previdencia",
     'Informações sobre planos de previdência privada (VGBL/PGBL) e seguros de vida detidos por {FIRST} e esposa, no Brasil e no exterior, conforme aplicável.'),
]


def _set_font(run, name: str = "Calibri", size: int = 11,
              bold: bool = False, color: tuple = (0, 0, 0)) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r, g, b = color
    run.font.color.rgb = RGBColor(r, g, b)


def _add_para(doc, text: str = "", *, align=None, space_after: int = 6,
              font_size: int = 11, bold: bool = False,
              color: tuple = (0, 0, 0)):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        _set_font(run, size=font_size, bold=bold, color=color)
    return p


def build_checklist(client_name: str, file_status: dict, output_path: str) -> str:
    """
    Gera o checklist preenchido.
    """
    client_name = (client_name or "Cliente").strip()
    name_upper = client_name.upper()
    first = name_upper.split()[0] if name_upper else "CLIENTE"
    file_status = file_status or {}

    doc = Document()

    # ── Margens ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    # ── Header com nome do escritório ────────────────────────────────
    _add_para(doc, "HUMBERTO SANCHES",
              align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=2, font_size=18, bold=True,
              color=(40, 57, 68))
    _add_para(doc, "+ ASSOCIADOS",
              align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=20, font_size=10, bold=True,
              color=(142, 149, 155))

    # ── Título do documento ──────────────────────────────────────────
    _add_para(doc, name_upper,
              align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=6, font_size=12, bold=True,
              color=(0, 0, 0))
    _add_para(doc, "PLANEJAMENTO PATRIMONIAL E SUCESSÓRIO",
              align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=2, font_size=11, bold=False,
              color=(0, 0, 0))
    _add_para(doc, "LISTA INICIAL: INFORMAÇÕES / DOCUMENTAÇÃO DE SUPORTE",
              align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=24, font_size=11, bold=False,
              color=(0, 0, 0))

    # ── Lista numerada com status ────────────────────────────────────
    for idx, (key, template) in enumerate(_ITEMS, start=1):
        is_done = bool(file_status.get(key, False))
        text = template.replace("{NAME}", name_upper).replace("{FIRST}", first)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)

        # Numeração
        num_run = p.add_run(f"{idx}.\t")
        _set_font(num_run, size=11)

        # Texto do item
        text_run = p.add_run(text)
        _set_font(text_run, size=11, color=(0, 0, 0))

        # Status (verde/vermelho, bold)
        if is_done:
            status_text = "  — Concluído"
            status_color = (46, 139, 87)  # verde
        else:
            status_text = "  — Pendente"
            status_color = (200, 16, 46)  # vermelho

        status_run = p.add_run(status_text)
        _set_font(status_run, size=11, bold=True, color=status_color)

    # ── Footer (rodapé) ──────────────────────────────────────────────
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    foot.paragraph_format.space_before = Pt(20)
    for line in [
        "SP +55 (11) 4858-7985",
        "hsanches@hsanches.com",
        "www.hsanches.com",
        "Av. Brigadeiro Faria Lima, 3200",
        "Edifício Seculum II, 2º andar",
        "São Paulo, SP · CEP 01451-000",
    ]:
        r = foot.add_run(line + "\n")
        _set_font(r, size=9, color=(120, 120, 120))

    doc.save(output_path)
    return output_path
